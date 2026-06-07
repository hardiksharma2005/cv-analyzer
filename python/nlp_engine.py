"""
python/nlp_engine.py — Core NLP Processing Engine
SMARRTIF AI CV Analyzer

Handles all heavy NLP work: text extraction from PDF/DOCX, spaCy-based
section parsing, skill detection with PhraseMatcher, experience/education
extraction using NER, and TF-IDF relevance scoring.

Designed to be imported by cv_analyzer.py (Flask API) and scorer.py.
"""

import os
import re
import json
import string
import logging
from pathlib import Path
from typing import Optional

import fitz                          # PyMuPDF — fast PDF rendering
import spacy
from spacy.matcher import PhraseMatcher
import nltk
from nltk.stem import PorterStemmer, WordNetLemmatizer
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# NLTK data bootstrap — called once at module import time.
# These corpora are required for stemming, lemmatisation, and stopword removal.
# ---------------------------------------------------------------------------
def _ensure_nltk_data():
    required = [
        ("tokenizers/punkt",         "punkt"),
        ("tokenizers/punkt_tab",     "punkt_tab"),
        ("corpora/stopwords",        "stopwords"),
        ("corpora/wordnet",          "wordnet"),
    ]
    for path, name in required:
        try:
            nltk.data.find(path)
        except LookupError:
            logger.info(f"Downloading NLTK corpus: {name}")
            nltk.download(name, quiet=True)

_ensure_nltk_data()

# ---------------------------------------------------------------------------
# Role profiles — loaded once at module level so every CVNLPEngine instance
# shares the same in-memory dict (avoids repeated disk I/O).
# ---------------------------------------------------------------------------
_PROFILES_PATH = Path(__file__).parent.parent / "data" / "role_profiles.json"

def _load_role_profiles() -> dict:
    try:
        with open(_PROFILES_PATH, encoding="utf-8") as f:
            return json.load(f).get("roles", {})
    except FileNotFoundError:
        logger.warning("role_profiles.json not found — skill vocabulary will be empty")
        return {}

_ROLE_PROFILES: dict = _load_role_profiles()

# ---------------------------------------------------------------------------
# Master skill vocabulary — union of all required + nice-to-have skills across
# every role.  Used to build the spaCy PhraseMatcher patterns.
# ---------------------------------------------------------------------------
def _build_skill_vocabulary() -> list[str]:
    skills = set()
    for profile in _ROLE_PROFILES.values():
        skills.update(s.lower() for s in profile.get("required_skills", []))
        skills.update(s.lower() for s in profile.get("nice_to_have_skills", []))
    return sorted(skills)

_SKILL_VOCAB: list[str] = _build_skill_vocabulary()

# ---------------------------------------------------------------------------
# Degree level mapping — mirrors the JS scorer's degree scoring table.
# Numeric keys represent confidence that the candidate holds that degree or higher.
# ---------------------------------------------------------------------------
DEGREE_LEVELS = {
    "phd": 5, "ph.d": 5, "doctorate": 5, "doctoral": 5,
    "master": 4, "msc": 4, "mba": 4, "meng": 4, "m.s": 4,
    "bachelor": 3, "bsc": 3, "b.s": 3, "b.e": 3, "b.tech": 3, "be": 3,
    "associate": 2,
    "diploma": 1, "certificate": 1,
}

# Action verbs commonly used in achievement-oriented CV writing.
# Used by analyze_writing_quality() to measure impact language density.
ACTION_VERBS = {
    "led", "built", "developed", "designed", "implemented", "created", "launched",
    "delivered", "managed", "optimized", "architected", "established", "drove",
    "reduced", "increased", "improved", "automated", "deployed", "integrated",
    "collaborated", "mentored", "trained", "facilitated", "negotiated", "analysed",
    "analyzed", "researched", "evaluated", "coordinated", "supervised", "executed",
    "spearheaded", "pioneered", "transformed", "streamlined", "generated",
}


class CVNLPEngine:
    """
    Central NLP engine for CV analysis.

    Lifecycle:
        engine = CVNLPEngine()      # loads spaCy model + builds PhraseMatcher
        text   = engine.extract_text_from_pdf(path)
        parsed = engine.parse_cv(text)
        skills = engine.extract_skills_nlp(text)
    """

    def __init__(self):
        # ------------------------------------------------------------------
        # Load spaCy's small English model.
        # en_core_web_sm is ~12 MB and provides tokenisation, POS tagging,
        # dependency parsing, and named entity recognition (NER).
        # The larger en_core_web_lg (~560 MB) has better NER accuracy but is
        # overkill for CV parsing where domain-specific patterns dominate.
        # ------------------------------------------------------------------
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            logger.error(
                "spaCy model 'en_core_web_sm' not found. "
                "Run: python -m spacy download en_core_web_sm"
            )
            raise

        # NLTK helpers for morphological normalisation
        self.stemmer    = PorterStemmer()
        self.lemmatizer = WordNetLemmatizer()
        self.stop_words = set(stopwords.words("english"))

        # Build the PhraseMatcher from the shared skill vocabulary.
        # LOWER attribute means matching is case-insensitive.
        self._matcher = self._build_phrase_matcher()

        logger.info(
            f"CVNLPEngine ready — "
            f"{len(_SKILL_VOCAB)} skills in vocabulary, "
            f"spaCy model: en_core_web_sm"
        )

    # -----------------------------------------------------------------------
    # PhraseMatcher construction
    # -----------------------------------------------------------------------
    def _build_phrase_matcher(self) -> PhraseMatcher:
        """
        Create a spaCy PhraseMatcher for every skill in the vocabulary.

        PhraseMatcher is much faster than regex for multi-token patterns because
        it operates on spaCy's internal hash IDs rather than string comparison.
        Matching ~200 skill patterns against a 500-word CV takes < 1 ms.
        """
        matcher = PhraseMatcher(self.nlp.vocab, attr="LOWER")
        patterns = [self.nlp.make_doc(skill) for skill in _SKILL_VOCAB]
        matcher.add("SKILLS", patterns)
        return matcher

    # -----------------------------------------------------------------------
    # Text extraction — PDF
    # -----------------------------------------------------------------------
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract plain text from a PDF using PyMuPDF (fitz).

        PyMuPDF is ~10× faster than pdfplumber and handles encrypted/scanned PDFs
        more gracefully.  We iterate pages and use get_text("text") which applies
        heuristic word-spacing to produce readable output without HTML overhead.

        Args:
            file_path: Absolute path to the PDF file.

        Returns:
            Extracted text with page breaks as double newlines.
        """
        pages = []
        try:
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc):
                page_text = page.get_text("text")
                if page_text.strip():
                    pages.append(page_text)
            doc.close()
        except Exception as exc:
            logger.error(f"PDF extraction failed for {file_path}: {exc}")
            raise ValueError(f"Could not read PDF: {exc}") from exc

        full_text = "\n\n".join(pages)
        return self._clean_text(full_text)

    # -----------------------------------------------------------------------
    # Text extraction — DOCX
    # -----------------------------------------------------------------------
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file using python-docx.

        python-docx reads the Open XML structure directly, preserving paragraph
        order and table cell content.  We explicitly iterate tables as well because
        many CV templates place skills or experience in table cells that plain
        paragraph iteration misses.

        Args:
            file_path: Absolute path to the DOCX file.

        Returns:
            Extracted text with paragraph-level newlines.
        """
        from docx import Document as DocxDocument
        try:
            doc = DocxDocument(file_path)
        except Exception as exc:
            raise ValueError(f"Could not read DOCX: {exc}") from exc

        lines = []

        # Standard paragraphs (headings, body text)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)

        # Table cells — common in modern CV templates
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in lines:
                        lines.append(cell_text)

        return self._clean_text("\n".join(lines))

    # -----------------------------------------------------------------------
    # Text cleaning
    # -----------------------------------------------------------------------
    def _clean_text(self, text: str) -> str:
        """
        Normalise whitespace and remove non-printable characters.
        Preserves newlines (important for section detection) while collapsing
        repeated blank lines to a single blank line.
        """
        # Remove null bytes and other control characters (except \n and \t)
        text = re.sub(r"[^\x09\x0a\x20-\x7e\x80-\xff]", " ", text)
        # Collapse runs of spaces/tabs to a single space
        text = re.sub(r"[ \t]+", " ", text)
        # Collapse runs of blank lines to two newlines (one blank line)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    # -----------------------------------------------------------------------
    # Full CV parsing
    # -----------------------------------------------------------------------
    def parse_cv(self, text: str) -> dict:
        """
        Main parsing entry point.

        Runs the full spaCy NLP pipeline on the CV text and delegates to
        specialised extractors for each section.  The pipeline uses:
          - Tokenisation + sentence boundary detection
          - POS tagging (used for action verb detection)
          - NER (used for DATE, ORG, GPE entities)

        Returns a structured dict compatible with the sessionStorage contract
        used by the JS frontend (minus raw_text which is stripped before storage).
        """
        # spaCy processes text as a single Doc object.
        # For very long CVs (> 10,000 chars) we increase the max_length limit.
        if len(text) > self.nlp.max_length:
            self.nlp.max_length = len(text) + 1000

        doc = self.nlp(text)

        skills_result    = self.extract_skills_nlp(text)
        experience_result = self.extract_experience_nlp(text, doc)
        education_result  = self.extract_education_nlp(text, doc)
        keywords          = self.extract_keywords(text, top_n=20)
        writing_quality   = self.analyze_writing_quality(text, doc)
        completeness      = self._calculate_profile_completeness(
            text, skills_result, experience_result, education_result
        )

        return {
            "skills": {
                "detected":    [s["skill"] for s in skills_result],
                "with_scores": skills_result,
            },
            "experience": experience_result,
            "education":  education_result,
            "keywords":   keywords,
            "writing_quality": writing_quality,
            "profile_completeness": completeness,
            "sections_detected": self._detect_sections(text),
            "contact_info": self._extract_contact_info(text),
            "word_count":   len(text.split()),
        }

    # -----------------------------------------------------------------------
    # Skill extraction with PhraseMatcher + stemming fallback
    # -----------------------------------------------------------------------
    def extract_skills_nlp(self, text: str) -> list[dict]:
        """
        Detect skills using spaCy's PhraseMatcher (primary) and NLTK stemming
        (secondary fallback for morphological variants).

        Primary path — PhraseMatcher:
            Matches exact and case-insensitive multi-word phrases from the
            vocabulary.  Returns span boundaries so we can deduplicate overlaps.

        Secondary path — NLTK stemming:
            For each vocabulary skill, stem every token and compare against
            stemmed CV tokens.  This catches variants like:
              "pythonic" → stem "python" → matches "python" skill
              "containerising" → stem "contain" → partial match for Docker context
            Stemmed matches receive a lower confidence score (0.6) to signal
            they are approximations.

        Returns:
            List of dicts: {skill, confidence, source}
            Sorted by confidence descending, deduplicated by skill name.
        """
        doc        = self.nlp(text)
        found      = {}   # skill_lower → {skill, confidence, source}

        # ── Primary: PhraseMatcher ──
        matches = self._matcher(doc)
        # Each match is (match_id, start, end) token indices
        for _match_id, start, end in matches:
            span      = doc[start:end]
            skill_key = span.text.lower()
            if skill_key not in found:
                found[skill_key] = {
                    "skill":      span.text,
                    "confidence": 0.95,  # direct phrase match → very high confidence
                    "source":     "phrase_match",
                }

        # ── Secondary: Stem-based fallback ──
        # Tokenise CV text and stem every content word (non-stopword, alphabetic)
        cv_tokens = text.lower().split()
        stemmed_cv_tokens = {
            self.stemmer.stem(t)
            for t in cv_tokens
            if t.isalpha() and t not in self.stop_words
        }

        for skill in _SKILL_VOCAB:
            if skill in found:
                continue  # already detected via phrase match
            # Stem each word of the skill phrase
            skill_words   = skill.split()
            stemmed_skill = {self.stemmer.stem(w) for w in skill_words if w.isalpha()}
            # Consider a match when ALL skill stems appear in the CV's stem set
            if stemmed_skill and stemmed_skill.issubset(stemmed_cv_tokens):
                found[skill] = {
                    "skill":      skill.title(),
                    "confidence": 0.60,
                    "source":     "stem_match",
                }

        # Sort by confidence (high → low) then alphabetically
        return sorted(found.values(), key=lambda x: (-x["confidence"], x["skill"]))

    # -----------------------------------------------------------------------
    # Experience extraction
    # -----------------------------------------------------------------------
    def extract_experience_nlp(self, text: str, doc=None) -> dict:
        """
        Extract work experience details using two complementary strategies.

        Strategy 1 — Explicit year statements (highest confidence):
            Regex patterns: "5 years of experience", "3+ years", "over 2 years"

        Strategy 2 — spaCy DATE entities (medium confidence):
            Finds DATE named entities (e.g. "Jan 2019 – Mar 2022") and parses
            year spans using regex on the entity text.

        Strategy 3 — Bare year ranges (lowest confidence):
            Raw regex on the full text for patterns like "2018-2021", "2020–Present".

        Years from strategies 2+3 are summed only when no Strategy 1 result is found,
        capping at 40 to avoid obviously bad parses (e.g. a graduation year range
        being mistaken for work experience).

        Returns:
            {
                total_years: float,
                confidence: str,          # "high" | "medium" | "low"
                positions: list[str],     # job titles detected
                companies: list[str],     # employer names (ORG entities)
                date_ranges: list[str],   # raw date strings found
            }
        """
        if doc is None:
            doc = self.nlp(text)

        # ── Strategy 1: explicit "X years" patterns ──
        explicit_patterns = [
            r"(\d+(?:\.\d+)?)\s*\+?\s*years?\s+(?:of\s+)?(?:work\s+)?experience",
            r"experience\s+of\s+(\d+(?:\.\d+)?)\s*\+?\s*years?",
            r"over\s+(\d+)\s+years?",
            r"more\s+than\s+(\d+)\s+years?",
        ]
        for pattern in explicit_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                years = float(match.group(1))
                return {
                    "total_years":  years,
                    "confidence":   "high",
                    "positions":    self._extract_job_titles(text),
                    "companies":    self._extract_companies(doc),
                    "date_ranges":  [],
                }

        # ── Strategy 2: spaCy DATE entities ──
        date_ranges  = []
        current_year = 2025  # use a fixed reference year for reproducibility
        total_months = 0

        year_span_pattern = re.compile(
            r"((?:19|20)\d{2})\s*[-–—to]+\s*((?:19|20)\d{2}|present|current|now)",
            re.IGNORECASE,
        )

        for ent in doc.ents:
            if ent.label_ == "DATE":
                raw = ent.text
                m = year_span_pattern.search(raw)
                if m:
                    start_yr = int(m.group(1))
                    end_raw  = m.group(2).lower()
                    end_yr   = current_year if end_raw in ("present", "current", "now") else int(m.group(2))
                    duration = max(0, end_yr - start_yr)
                    if 0 < duration <= 15:  # sanity-check: no single job > 15 years
                        total_months += duration * 12
                        date_ranges.append(raw)

        # ── Strategy 3: fallback bare year ranges in plain text ──
        if not date_ranges:
            for m in year_span_pattern.finditer(text):
                start_yr = int(m.group(1))
                end_raw  = m.group(2).lower()
                end_yr   = current_year if end_raw in ("present", "current", "now") else int(m.group(2))
                duration = max(0, end_yr - start_yr)
                if 0 < duration <= 15:
                    total_months += duration * 12
                    date_ranges.append(m.group(0))

        total_years = min(round(total_months / 12, 1), 40)

        return {
            "total_years": total_years,
            "confidence":  "medium" if date_ranges else "low",
            "positions":   self._extract_job_titles(text),
            "companies":   self._extract_companies(doc),
            "date_ranges": date_ranges,
        }

    def _extract_job_titles(self, text: str) -> list[str]:
        """
        Heuristic job title extraction.
        Looks for lines/phrases where common title indicators appear.
        """
        title_keywords = [
            "engineer", "analyst", "scientist", "developer", "architect",
            "manager", "lead", "senior", "junior", "associate", "consultant",
            "intern", "director", "head of", "vp of",
        ]
        titles = []
        for line in text.splitlines():
            line_lower = line.lower().strip()
            if any(kw in line_lower for kw in title_keywords):
                # Avoid capturing full paragraphs — limit to short lines
                if 3 <= len(line.split()) <= 8:
                    titles.append(line.strip())
        return titles[:5]  # return top 5 to avoid noise

    def _extract_companies(self, doc) -> list[str]:
        """
        Extract ORG named entities as likely employer names.
        Filters out known false positives (university names handled by education extractor).
        """
        university_signals = {"university", "college", "institute", "school", "academy"}
        companies = []
        seen = set()
        for ent in doc.ents:
            if ent.label_ == "ORG":
                name = ent.text.strip()
                if name.lower() not in seen and len(name) > 2:
                    # Skip if it looks like an educational institution
                    if not any(s in name.lower() for s in university_signals):
                        companies.append(name)
                        seen.add(name.lower())
        return companies[:10]

    # -----------------------------------------------------------------------
    # Education extraction
    # -----------------------------------------------------------------------
    def extract_education_nlp(self, text: str, doc=None) -> dict:
        """
        Extract education details by combining spaCy NER (ORG entities for
        institution names, DATE entities for graduation years) with regex
        pattern matching for degree keywords.

        Returns:
            {
                degree_level: int,          # 0-5 matching JS scorer table
                degree_name: str,
                institution: str,
                graduation_year: int | None,
                field_keywords: list[str],
            }
        """
        if doc is None:
            doc = self.nlp(text)

        # ── Degree level detection ──
        best_level = 0
        best_degree = "Not detected"
        text_lower = text.lower()

        for keyword, level in DEGREE_LEVELS.items():
            if re.search(r"\b" + re.escape(keyword) + r"\b", text_lower):
                if level > best_level:
                    best_level  = level
                    best_degree = keyword.upper()

        # ── Institution detection (ORG entities with university signals) ──
        institution = "Not detected"
        uni_signals = {"university", "college", "institute", "school", "academy", "polytechnic"}
        for ent in doc.ents:
            if ent.label_ == "ORG":
                if any(s in ent.text.lower() for s in uni_signals):
                    institution = ent.text.strip()
                    break

        # ── Graduation year ──
        grad_year = None
        year_pattern = re.compile(r"\b(19[89]\d|20[0-2]\d)\b")
        years_found = [int(m.group()) for m in year_pattern.finditer(text)]
        # Heuristic: graduation year is typically the largest year ≤ current year
        valid_years = [y for y in years_found if 1980 <= y <= 2025]
        if valid_years:
            grad_year = max(valid_years)

        # ── Field keyword matching ──
        all_field_keywords = set()
        for profile in _ROLE_PROFILES.values():
            all_field_keywords.update(kw.lower() for kw in profile.get("education_keywords", []))

        detected_fields = [
            kw for kw in all_field_keywords
            if re.search(r"\b" + re.escape(kw) + r"\b", text_lower)
        ]

        return {
            "degree_level":    best_level,
            "degree_name":     best_degree,
            "institution":     institution,
            "graduation_year": grad_year,
            "field_keywords":  detected_fields,
        }

    # -----------------------------------------------------------------------
    # TF-IDF role relevance
    # -----------------------------------------------------------------------
    def calculate_tfidf_relevance(self, cv_text: str, role: str) -> float:
        """
        Compute cosine similarity between the CV text and a synthetic role
        description built from the role's required skills, education keywords,
        and certification names.

        TF-IDF vectorisation weights rare terms higher than common ones, so a
        CV that mentions obscure but role-critical skills (e.g. "MLflow" for a
        Data Scientist) is scored higher than one that only mentions generic
        terms ("data", "analysis").

        Args:
            cv_text: Full CV text.
            role: One of the keys in role_profiles.json.

        Returns:
            Similarity score 0–100 (rounded to 1 decimal).
        """
        profile = _ROLE_PROFILES.get(role)
        if not profile:
            return 0.0

        # Build a synthetic "role document" from profile terms
        role_terms = (
            profile.get("required_skills", [])
            + profile.get("nice_to_have_skills", [])
            + profile.get("education_keywords", [])
            + profile.get("key_certifications", [])
        )
        role_doc = " ".join(role_terms).lower()

        # TF-IDF on two documents: CV text vs role description
        try:
            vectorizer = TfidfVectorizer(
                stop_words="english",
                ngram_range=(1, 2),   # unigrams + bigrams to capture "machine learning"
                min_df=1,
                sublinear_tf=True,    # log-scale TF to reduce impact of very frequent terms
            )
            matrix = vectorizer.fit_transform([cv_text.lower(), role_doc])
            similarity = cosine_similarity(matrix[0:1], matrix[1:2])[0][0]
            return round(float(similarity) * 100, 1)
        except Exception as exc:
            logger.warning(f"TF-IDF similarity failed: {exc}")
            return 0.0

    # -----------------------------------------------------------------------
    # Keyword extraction
    # -----------------------------------------------------------------------
    def extract_keywords(self, text: str, top_n: int = 20) -> list[dict]:
        """
        Extract the most important keywords from a CV using TF-IDF on
        character n-grams combined with word n-grams.

        We treat the full CV as a single document and rank terms by their
        TF-IDF score.  Stopwords and single-character tokens are filtered.

        Returns:
            List of {keyword, score} dicts, sorted by score descending.
        """
        # Clean text: remove punctuation, lowercase
        cleaned = re.sub(r"[^\w\s]", " ", text.lower())

        try:
            vectorizer = TfidfVectorizer(
                stop_words="english",
                ngram_range=(1, 2),
                min_df=1,
                max_features=200,
            )
            # fit_transform needs an iterable — single document is valid
            tfidf_matrix = vectorizer.fit_transform([cleaned])
            feature_names = vectorizer.get_feature_names_out()
            scores        = tfidf_matrix.toarray()[0]

            # Zip feature names with scores and sort
            keyword_scores = sorted(
                zip(feature_names, scores),
                key=lambda x: x[1],
                reverse=True,
            )

            # Filter out tokens that are purely numeric or < 3 chars
            results = []
            for kw, score in keyword_scores:
                if len(kw) >= 3 and not kw.replace(" ", "").isnumeric():
                    results.append({"keyword": kw, "score": round(float(score), 4)})
                if len(results) >= top_n:
                    break

            return results
        except Exception as exc:
            logger.warning(f"Keyword extraction failed: {exc}")
            return []

    # -----------------------------------------------------------------------
    # Writing quality analysis
    # -----------------------------------------------------------------------
    def analyze_writing_quality(self, text: str, doc=None) -> dict:
        """
        Compute lightweight NLP metrics that proxy for CV writing quality.

        Metrics:
          - avg_sentence_length: mean tokens per sentence (ideal: 12–20)
          - vocabulary_richness: unique words / total words (higher = better)
          - action_verb_density: fraction of sentences beginning with action verbs
          - quantification_score: fraction of sentences containing a number
          - quality_score: composite 0–100 from the above metrics

        These metrics are used as a bonus signal in ATS scoring, not as primary
        scoring dimensions.  A well-written CV with strong action verbs and
        quantified achievements scores higher.
        """
        if doc is None:
            doc = self.nlp(text)

        sentences = list(doc.sents)
        if not sentences:
            return {"quality_score": 0, "details": {}}

        # Average sentence length (in tokens, excluding whitespace tokens)
        sent_lengths = [
            len([t for t in s if not t.is_space])
            for s in sentences
        ]
        avg_len = sum(sent_lengths) / len(sent_lengths) if sent_lengths else 0

        # Vocabulary richness
        all_words   = [t.lower_ for t in doc if t.is_alpha and not t.is_stop]
        vocab_ratio = len(set(all_words)) / len(all_words) if all_words else 0

        # Action verb density
        action_sents = 0
        for sent in sentences:
            # First non-space token that is a VERB or propn in lemma
            first_verb = next(
                (t for t in sent if not t.is_space and t.is_alpha),
                None,
            )
            if first_verb and first_verb.lemma_.lower() in ACTION_VERBS:
                action_sents += 1
        action_density = action_sents / len(sentences)

        # Quantification score — sentences containing at least one numeral
        number_pattern = re.compile(r"\b\d[\d,\.]*[%x]?\b")
        quant_sents    = sum(
            1 for s in sentences if number_pattern.search(s.text)
        )
        quant_score    = quant_sents / len(sentences)

        # Composite quality score (0–100)
        # Sentence length: ideal 12–20 tokens → linear penalty outside range
        len_score = max(0, 1 - abs(avg_len - 16) / 20)
        quality   = round(
            (len_score * 20 + vocab_ratio * 30 + action_density * 30 + quant_score * 20),
            1,
        )

        return {
            "quality_score":       min(100, quality),
            "details": {
                "avg_sentence_length":  round(avg_len, 1),
                "vocabulary_richness":  round(vocab_ratio * 100, 1),
                "action_verb_density":  round(action_density * 100, 1),
                "quantification_score": round(quant_score * 100, 1),
            },
        }

    # -----------------------------------------------------------------------
    # Section detection
    # -----------------------------------------------------------------------
    def _detect_sections(self, text: str) -> list[str]:
        """
        Identify standard CV section headers using regex.
        Section presence is used for ATS compatibility scoring.
        """
        section_patterns = {
            "summary":        r"\b(summary|profile|objective|about)\b",
            "experience":     r"\b(experience|employment|work history|career)\b",
            "education":      r"\b(education|academic|qualification|degree)\b",
            "skills":         r"\b(skills|competencies|technologies|expertise)\b",
            "projects":       r"\b(projects|portfolio|work samples)\b",
            "certifications": r"\b(certification|certificate|credential|licence)\b",
            "contact":        r"\b(contact|email|phone|linkedin|github)\b",
            "achievements":   r"\b(achievement|award|recognition|honour|honor)\b",
        }
        text_lower = text.lower()
        detected   = []
        for section, pattern in section_patterns.items():
            if re.search(pattern, text_lower):
                detected.append(section)
        return detected

    # -----------------------------------------------------------------------
    # Contact info extraction
    # -----------------------------------------------------------------------
    def _extract_contact_info(self, text: str) -> dict:
        """
        Detect contact information completeness.
        Returns presence flags only — no actual data is stored.
        """
        return {
            "has_email":    bool(re.search(r"[\w.+-]+@[\w-]+\.[a-z]{2,}", text, re.IGNORECASE)),
            "has_phone":    bool(re.search(r"[\+\d][\d\s\-\(\)]{7,}\d", text)),
            "has_linkedin": bool(re.search(r"linkedin\.com", text, re.IGNORECASE)),
            "has_github":   bool(re.search(r"github\.com", text, re.IGNORECASE)),
        }

    # -----------------------------------------------------------------------
    # Profile completeness (mirrors JS parser logic, extended with NLP signals)
    # -----------------------------------------------------------------------
    def _calculate_profile_completeness(
        self,
        text: str,
        skills: list,
        experience: dict,
        education: dict,
    ) -> dict:
        """
        Calculate a completeness score (0–100) based on the presence and richness
        of each CV section, mirroring the JS parser's logic.

        Weights (sum to 100):
            experience     → 30
            skills         → 25
            education      → 15
            contact        → 10
            summary        → 10
            projects       → 5
            certifications → 5
        """
        contact = self._extract_contact_info(text)
        text_lower = text.lower()

        # Experience richness: tiered by years detected
        yrs = experience.get("total_years", 0)
        exp_score = (
            100 if yrs >= 5 else
            85  if yrs >= 3 else
            60  if yrs >= 1 else
            30  if yrs > 0  else
            0
        )

        # Skills richness: tiered by number detected
        n_skills = len(skills)
        skill_score = (
            100 if n_skills >= 15 else
            85  if n_skills >= 10 else
            60  if n_skills >= 5  else
            30  if n_skills >= 1  else
            0
        )

        # Education
        edu_score = min(100, education.get("degree_level", 0) * 20)

        # Contact: each present field adds 25 points, max 100
        contact_score = min(100, sum(25 for v in contact.values() if v))

        # Summary / objective
        has_summary = bool(re.search(
            r"\b(summary|profile|objective|about me)\b", text_lower
        ))
        summary_score = 100 if has_summary else 0

        # Projects
        has_projects = bool(re.search(r"\b(project|portfolio)\b", text_lower))
        project_score = 100 if has_projects else 0

        # Certifications
        has_certs = bool(re.search(
            r"\b(certification|certificate|certified|credential)\b", text_lower
        ))
        cert_score = 100 if has_certs else 0

        weighted = (
            exp_score    * 0.30 +
            skill_score  * 0.25 +
            edu_score    * 0.15 +
            contact_score * 0.10 +
            summary_score * 0.10 +
            project_score * 0.05 +
            cert_score   * 0.05
        )

        return {
            "score":   round(weighted, 1),
            "breakdown": {
                "experience":     exp_score,
                "skills":         skill_score,
                "education":      edu_score,
                "contact":        contact_score,
                "summary":        summary_score,
                "projects":       project_score,
                "certifications": cert_score,
            },
        }
